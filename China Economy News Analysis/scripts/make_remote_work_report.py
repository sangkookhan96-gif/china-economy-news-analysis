#!/usr/bin/env python3
"""외교부 제출 '일일 원격근무 실적보고서(한상국).docx' 자동 생성.

규칙(메모리 mofa-daily-report 참조):
- 평일만(주말·공휴일 제외). 보고일 = 대상일.
- 업무 실적 = 그날 공개된 플랫폼 주요 뉴스 제목 2건을 "(제목) 동향에 대한 분석"으로 기재(오전 1, 오후 1).
- 양식: 기존 .docx 템플릿 복제 후 날짜·오전·오후 셀만 교체(서식 보존).
- 안전: 동일 파일이 이미 있으면 덮어쓰지 않음(--force 시에만). 사용자 수동 작성본 보호.

사용:
  python3 scripts/make_remote_work_report.py            # 오늘
  python3 scripts/make_remote_work_report.py 2026-06-18 # 특정일
  python3 scripts/make_remote_work_report.py --force    # 기존 파일 덮어쓰기 허용
"""
import sys
import os
import re
import shutil
import datetime as dt

sys.path.insert(0, "/home/jeozeohan/vibe_temp/China Economy News Analysis")
from src.database.models import get_connection

REPORT_DIR = "/home/jeozeohan/Documents/일일근무보고서"
TEMPLATE = os.path.join(REPORT_DIR, "0617 일일 원격근무 실적보고서(한상국).docx")
WEEKDAY_KO = ["월", "화", "수", "목", "금", "토", "일"]

# 한국 공휴일 2026 (주요 법정공휴일·대체공휴일)
HOLIDAYS_2026 = {
    "2026-01-01",  # 신정
    "2026-02-16", "2026-02-17", "2026-02-18",  # 설 연휴
    "2026-03-01", "2026-03-02",  # 삼일절(일)+대체
    "2026-05-05",  # 어린이날
    "2026-05-24", "2026-05-25",  # 부처님오신날(일)+대체
    "2026-06-06",  # 현충일
    "2026-08-15", "2026-08-17",  # 광복절(토)+대체
    "2026-09-24", "2026-09-25", "2026-09-26",  # 추석 연휴
    "2026-10-03", "2026-10-05",  # 개천절(토)+대체
    "2026-10-09",  # 한글날
    "2026-12-25",  # 성탄절
}

# 크롤링 잔재·낚시성·불완전 헤드라인 배제용
_BAD_RE = re.compile(r"\d{4,}|\d+\s*(분|시간)\s*전|[?？]$|VS|일문|看懂")
_INCOMPLETE_TAIL = ("의", "을", "를", "와", "과", "에", "에서", "으로", "로", "은", "는",
                    "이", "가", "도", "및", "등", "한", "할", "하는", "되는", "위한")


def _is_clean(h):
    h = (h or "").strip()
    if len(h) < 6:
        return False
    if _BAD_RE.search(h):
        return False
    if any(h.endswith(t) for t in _INCOMPLETE_TAIL):
        return False
    return True


def pick_headlines(date_str):
    """그날 공개 뉴스에서 중요도 상위 정상 헤드라인 2건."""
    c = get_connection().cursor()
    c.execute("""SELECT n.card_headline h, n.importance_score imp
                 FROM news n JOIN cni_summaries cs ON cs.news_id = n.id
                 WHERE cs.published_at IS NOT NULL AND date(cs.published_at)=?""", (date_str,))
    seen, cands = set(), []
    for r in c.fetchall():
        h = (r["h"] or "").strip()
        if not _is_clean(h) or h in seen:
            continue
        seen.add(h)
        has_num = 1 if re.search(r"\d|%|％", h) else 0   # 수치 포함(구체 사실) 우선
        cands.append((has_num, r["imp"] or 0.0, h))
    # 수치 포함 → 중요도 순. 동률은 안정 정렬.
    cands.sort(key=lambda x: (x[0], x[1]), reverse=True)
    return [c[2] for c in cands[:2]]


def _set_cell(cell, text):
    """병합셀 텍스트 교체 — 첫 run 서식 보존, 나머지 run 제거."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _distinct_cells(row):
    out, last = [], None
    for cell in row.cells:
        if cell._tc is not last:
            out.append(cell)
            last = cell._tc
    return out


def build(date_obj, force=False):
    ds = date_obj.strftime("%Y-%m-%d")
    if date_obj.weekday() >= 5:
        print(f"[skip] {ds} 주말 — 보고서 미작성")
        return None
    if ds in HOLIDAYS_2026:
        print(f"[skip] {ds} 공휴일 — 보고서 미작성")
        return None

    mmdd = date_obj.strftime("%m%d")
    out_path = os.path.join(REPORT_DIR, f"{mmdd} 일일 원격근무 실적보고서(한상국).docx")
    if os.path.exists(out_path) and not force:
        print(f"[skip] 이미 존재(보호): {out_path}  (--force로 덮어쓰기)")
        return None

    picks = pick_headlines(ds)
    if len(picks) < 2:
        print(f"[warn] {ds} 정상 공개 헤드라인 부족({len(picks)}건) — 작성 보류")
        return None

    am = f"{picks[0]} 동향에 대한 분석"
    pm = f"{picks[1]} 동향에 대한 분석"
    date_cell_txt = f"‘26.{date_obj.month}.{date_obj.day}.({WEEKDAY_KO[date_obj.weekday()]})"

    shutil.copyfile(TEMPLATE, out_path)
    import docx
    d = docx.Document(out_path)
    t = d.tables[1]
    _set_cell(_distinct_cells(t.rows[6])[1], date_cell_txt)   # 근무일자
    _set_cell(_distinct_cells(t.rows[9])[1], am)              # 오전
    _set_cell(_distinct_cells(t.rows[13])[1], pm)             # 오후
    d.save(out_path)

    print(f"[done] {out_path}")
    print(f"  일자: {date_cell_txt}")
    print(f"  오전: {am}")
    print(f"  오후: {pm}")
    return out_path


def main():
    args = [a for a in sys.argv[1:] if a != "--force"]
    force = "--force" in sys.argv
    if args:
        date_obj = dt.datetime.strptime(args[0], "%Y-%m-%d").date()
    else:
        date_obj = dt.date.today()
    build(date_obj, force=force)


if __name__ == "__main__":
    main()
